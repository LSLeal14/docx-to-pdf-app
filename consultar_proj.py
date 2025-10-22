import streamlit as st
import os
import tempfile
import shutil
import re
import subprocess
from pathlib import Path
import firebase_admin
from firebase_admin import credentials, firestore
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

# Funções de processamento de tabelas
from processamento import (
    gerar_tabela_percentual, 
    gerar_tabela_previsto_realizado, 
    gerar_tabela_previsto_realizado_mes, 
    gerar_tabela_contratual, 
    gerar_tabela_previsto_realizado_acumulado
)

# --- NOVO: Importar funções de gráfico ---
from data_gen.graphs import (
    gerar_curva_s,
    gerar_grafico_aderencia
)
# ----------------------------------------


def get_downloads_folder():
    """Retorna o caminho para a pasta de Downloads do usuário."""
    home = Path.home()
    return home / "Downloads"

def extrair_campos(doc):
    """Extrai todos os placeholders {{campo}} de um documento Word."""
    campos = set()
    for p in doc.paragraphs:
        campos.update(re.findall(r"\{\{(.*?)\}\}", p.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    campos.update(re.findall(r"\{\{(.*?)\}\}", p.text))
    return list(campos)

def preencher_campos(doc, dados):
    """
    Substitui placeholders {{chave}} no documento por:
      - Texto (str/int/float)
      - Imagem (caminho .png/.jpg/.jpeg)
      - Tabela (list[dict] ou pandas.DataFrame) com grid.
    """
    def is_image(v):
        # --- MODIFICAÇÃO: Garantir que Path ou str funcionem ---
        return isinstance(v, (str, Path)) and str(v).lower().endswith((".png", ".jpg", ".jpeg"))
        # ------------------------------------------------------

    def is_table(v):
        if isinstance(v, pd.DataFrame):
            return True
        if isinstance(v, list) and v and all(isinstance(r, dict) for r in v):
            return True
        return False

    def normalize_table(v):
        if isinstance(v, pd.DataFrame):
            return v.to_dict(orient="records")
        return v

    def replace_text_in_paragraph(p, mapping_text):
        original_text = p.text
        for k, v in mapping_text.items():
            placeholder = f"{{{{{k}}}}}"
            if placeholder in p.text:
                # Substitui o placeholder pelo valor, mantendo o resto do texto
                # --- MODIFICAÇÃO: Converte 'v' para str explicitamente ---
                p.text = p.text.replace(placeholder, str(v) if v is not None else "")
                # -----------------------------------------------------------

    def paragraph_only_placeholder(p, placeholder):
        return p.text.strip() == placeholder

    def insert_image_at_paragraph(p, img_path):
        for r in p.runs:
            r.text = ""
        # --- MODIFICAÇÃO: Garantir que img_path é string ---
        p.add_run().add_picture(str(img_path), width=Inches(6.0)) # Aumentei a largura
        # -----------------------------------------------------

    def insert_table_after_paragraph(doc, p, records):
        """Cria uma tabela com grid e cabeçalho em negrito."""
        if not records:
            p.text = "" # Limpa o placeholder se a tabela estiver vazia
            return
        
        cols = list(records[0].keys())
        
        table = doc.add_table(rows=1 + len(records), cols=len(cols), style='Table Grid')
        table.autofit = True # Ajusta colunas ao conteúdo

        hdr_cells = table.rows[0].cells
        for j, c in enumerate(cols):
            run = hdr_cells[j].paragraphs[0].add_run(str(c))
            run.bold = True
            run.font.size = Pt(10)

        for i, row_data in enumerate(records, start=1):
            row_cells = table.rows[i].cells
            for j, c in enumerate(cols):
                cell_value = row_data.get(c)
                run = row_cells[j].paragraphs[0].add_run("" if cell_value is None else str(cell_value))
                run.font.size = Pt(10)
                
        p._element.addnext(table._element)
        p.text = "" # Limpa o placeholder

    dados_texto, dados_imagem, dados_tabela = {}, {}, {}
    for k, v in dados.items():
        if is_image(v):
            dados_imagem[k] = v
        elif is_table(v):
            dados_tabela[k] = normalize_table(v)
        else:
            dados_texto[k] = v

    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for p in all_paragraphs:
        # --- MODIFICAÇÃO: Inserir tabelas e imagens primeiro ---
        
        # Lógica de Tabela (substitui o placeholder)
        for k, rows in dados_tabela.items():
            ph = f"{{{{{k}}}}}"
            if ph in p.text:
                if paragraph_only_placeholder(p, ph):
                    insert_table_after_paragraph(doc, p, rows)
                else:
                    # Se houver texto junto, apenas limpa (a tabela será inserida depois)
                    p.text = p.text.replace(ph, "")
                    insert_table_after_paragraph(doc, p, rows)

        # Lógica de Imagem (substitui o placeholder)
        for k, img in dados_imagem.items():
            ph = f"{{{{{k}}}}}"
            if ph in p.text:
                if paragraph_only_placeholder(p, ph):
                    insert_image_at_paragraph(p, img)
                else:
                    # Se houver mais texto, a imagem é inserida inline
                    p.text = p.text.replace(ph, "")
                    p.add_run().add_picture(str(img), width=Inches(6.0))

        # Lógica de Texto (substitui o que sobrou)
        replace_text_in_paragraph(p, dados_texto)
        # ----------------------------------------------------

def converter_para_pdf(caminho_docx):
    """Converte um arquivo .docx para .pdf usando LibreOffice."""
    downloads_dir = get_downloads_folder()
    os.makedirs(downloads_dir, exist_ok=True)
    comando = [
        "soffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(downloads_dir), caminho_docx
    ]
    result = subprocess.run(comando, capture_output=True, text=True, check=False)
    if result.returncode != 0:
        raise RuntimeError(f"Erro na conversão para PDF:\n{result.stderr}")

    base_name = Path(caminho_docx).stem
    pdf_path = downloads_dir / f"{base_name}.pdf"
    
    if not pdf_path.exists():
         raise FileNotFoundError(f"Arquivo PDF esperado não foi encontrado em: {pdf_path}")
    
    return str(pdf_path)

# ==== App principal ====
def main():
    st.set_page_config(layout="wide")
    st.title("Consulta de Projetos e Geração de PDF")

    # Inicialização do Firebase (evita reinicializar)
    if not firebase_admin._apps:
        try:
            FIREBASE_KEY_PATH = os.getenv("FIREBASE_KEY_PATH", "app/firebase_key.json")
            cred = credentials.Certificate(FIREBASE_KEY_PATH)
            firebase_admin.initialize_app(cred)
        except Exception as e:
            st.error(f"Erro ao inicializar Firebase: {e}")
            st.stop()

    db = firestore.client()

    campos = {
        "N° do Contrato": "n_contrato", "Período de Vigência": "periodo_vigencia",
        "N° da OS/OFB/NE": "n_os", "Objeto": "objeto",
        "Valor dos Bens/Serviços Recebidos": "valor_bens_receb",
        "Contratante": "contratante", "Contratada": "contratada"
    }

    campo_escolhido = st.selectbox("Selecione o campo para buscar:", list(campos.keys()))
    termo_busca = st.text_input("Digite o termo para busca:")

    projetos_ref = db.collection("projetos")
    
    st.subheader("Projetos encontrados:")
    campo_firebase = campos[campo_escolhido]
    
    query = projetos_ref.stream()
    resultados = []
    for doc in query:
        data = doc.to_dict()
        if termo_busca.lower() in str(data.get(campo_firebase, "")).lower():
            resultados.append((doc.id, data))

    if resultados:
        for doc_id, data in resultados:
            st.markdown(f"---")
            df_info = pd.DataFrame({
                "Item": list(campos.keys()),
                "Info": [str(data.get(campos[campo], "")) for campo in campos]
            })
            st.table(df_info)

            if st.button(f"Gerar PDF para o Projeto", key=f"gerar_pdf_{doc_id}"):
                
                # --- NOVO: Lista para guardar caminhos de imagens temporárias ---
                temp_image_paths = []
                # --------------------------------------------------------------
                
                with st.spinner("Gerando tabelas e gráficos..."):
                    try:
                        # --- GERAÇÃO DE TABELAS ---
                        tabela_1_df = gerar_tabela_percentual(db, doc_id)
                        if tabela_1_df is None:
                            st.error("Falha ao gerar a tabela 1 (Percentual).")
                            continue

                        tabela_2_df = gerar_tabela_previsto_realizado(db, doc_id)
                        if tabela_2_df is None:
                            st.error("Falha ao gerar a tabela 2 (Previsto x Realizado).")
                            continue

                        tabela_3_df = gerar_tabela_previsto_realizado_mes(db, doc_id)
                        if tabela_3_df is None:
                            st.error("Falha ao gerar a tabela 3 (Mês a Mês).")
                            continue

                        tabela_4_df = gerar_tabela_contratual(db, doc_id)
                        if tabela_4_df is None:
                            st.error("Falha ao gerar a tabela 4 (Contratual).")
                            continue

                        tabela_5_df = gerar_tabela_previsto_realizado_acumulado(db, doc_id)
                        if tabela_5_df is None:
                            st.error("Falha ao gerar a tabela 5 (Acumulado).")
                            continue
                        
                        # --- NOVO: GERAÇÃO DE GRÁFICOS ---
                        # Gráfico para {{grafico_1}} e {{grafico_2}} (Aderência)
                        path_grafico_aderencia = gerar_grafico_aderencia(tabela_5_df)
                        if path_grafico_aderencia:
                            temp_image_paths.append(path_grafico_aderencia)
                        
                        # Gráfico para {{grafico_3}} e {{grafico_4}} (Curva S)
                        path_curva_s = gerar_curva_s(tabela_3_df)
                        if path_curva_s:
                            temp_image_paths.append(path_curva_s)
                        
                        # (Você pode adicionar aqui a geração do grafico_7 se criar a função)
                        # -----------------------------------

                        # 2. Adiciona tabelas e gráficos ao dicionário
                        dados_para_template = data.copy()
                        dados_para_template['table'] = tabela_1_df
                        dados_para_template['table_2'] = tabela_2_df
                        dados_para_template['table_3'] = tabela_3_df
                        dados_para_template['table_4'] = tabela_4_df
                        dados_para_template['table_5'] = tabela_5_df
                        # (Adicione table_6, table_7... se existirem)

                        # --- NOVO: Adiciona gráficos ---
                        # Usando os mesmos gráficos para os placeholders, conforme template
                        dados_para_template['grafico_1'] = path_grafico_aderencia
                        dados_para_template['grafico_2'] = path_grafico_aderencia
                        dados_para_template['grafico_3'] = path_curva_s
                        dados_para_template['grafico_4'] = path_curva_s
                        dados_para_template['grafico_5'] = path_curva_s # Placeholder existe, vamos preencher
                        dados_para_template['grafico_6'] = path_curva_s # Placeholder existe, vamos preencher
                        # dados_para_template['grafico_7'] = path_grafico_idp 
                        # ---------------------------------
                        
                        # Copia o template para um local temporário
                        caminho_template = "template/Template_ata_ebserh.docx"
                        if not os.path.exists(caminho_template):
                            st.error(f"Template não encontrado: {caminho_template}")
                            continue
                        
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                            shutil.copyfile(caminho_template, temp_docx.name)
                            
                            doc_obj = Document(temp_docx.name)
                            preencher_campos(doc_obj, dados_para_template)
                            doc_obj.save(temp_docx.name)
                            
                            pdf_path = converter_para_pdf(temp_docx.name)

                        with open(pdf_path, "rb") as pdf_file:
                            st.download_button(
                                label="✔️ Baixar PDF Pronto",
                                data=pdf_file,
                                file_name=f"projeto_{doc_id}.pdf",
                                mime="application/pdf",
                                key=f"download_{doc_id}"
                            )
                        st.success(f"PDF gerado com sucesso! Salvo em sua pasta de Downloads.")
                    
                    except Exception as e:
                        st.error(f"Erro ao gerar PDF: {e}")
                    
                    # --- NOVO: Bloco Finally para limpar imagens ---
                    finally:
                        for img_path in temp_image_paths:
                            try:
                                if img_path and os.path.exists(img_path):
                                    os.remove(img_path)
                            except Exception as e:
                                st.warning(f"Não foi possível remover o arquivo temporário {img_path}: {e}")
                    # ----------------------------------------------
    else:
        st.info("Nenhum projeto encontrado com os critérios de busca.")

if __name__ == "__main__":
    main()
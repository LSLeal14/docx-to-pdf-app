import streamlit as st
from docx import Document
import tempfile
import shutil
import os
import re
import subprocess
from pathlib import Path

def get_downloads_folder():
    home = Path.home()
    if os.name == 'nt':  # Windows
        downloads = home / "Downloads"
    else:  # macOS/Linux
        downloads = home / "Downloads"
    return downloads

def extrair_campos(doc):
    campos = set()
    for p in doc.paragraphs:
        campos.update(re.findall(r"\{\{(.*?)\}\}", p.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    campos.update(re.findall(r"\{\{(.*?)\}\}", p.text))
    return list(campos)

def preencher_campos(doc, dados):
    for p in doc.paragraphs:
        for k, v in dados.items():
            p.text = p.text.replace(f"{{{{{k}}}}}", v)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in dados.items():
                        p.text = p.text.replace(f"{{{{{k}}}}}", v)

def converter_para_pdf(caminho_docx):
    downloads_dir = get_downloads_folder()
    os.makedirs(downloads_dir, exist_ok=True)

    comando = [
        "soffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", str(downloads_dir),
        caminho_docx
    ]
    st.write(f"Executando comando: {' '.join(comando)}")
    result = subprocess.run(comando, capture_output=True, text=True)

    st.write("STDOUT:", result.stdout)
    st.write("STDERR:", result.stderr)

    if result.returncode != 0:
        raise RuntimeError(f"Erro na conversão para PDF:\n{result.stderr}")

    arquivos = os.listdir(downloads_dir)
    pdfs = [f for f in arquivos if f.lower().endswith(".pdf")]

    if not pdfs:
        raise FileNotFoundError(f"Nenhum arquivo PDF encontrado na pasta {downloads_dir}")

    pdfs = sorted(pdfs, key=lambda f: os.path.getmtime(os.path.join(downloads_dir, f)), reverse=True)
    caminho_pdf = os.path.join(downloads_dir, pdfs[0])

    return caminho_pdf

st.title("📄 Gerar PDF e salvar na pasta Downloads")

caminho_fixo = "template/Template_ata_ebserh.docx"

if not os.path.exists(caminho_fixo):
    st.error(f"Arquivo fixo não encontrado: {caminho_fixo}")
else:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        shutil.copyfile(caminho_fixo, tmp.name)
        caminho_docx_temp = tmp.name

    doc = Document(caminho_docx_temp)
    campos = extrair_campos(doc)

    if campos:
        st.success("Campos encontrados:")
        dados = {}
        for campo in campos:
            dados[campo] = st.text_input(campo)

        if st.button("Gerar PDF"):
            preencher_campos(doc, dados)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as docx_preenchido:
                doc.save(docx_preenchido.name)
                pdf_path = converter_para_pdf(docx_preenchido.name)

                with open(pdf_path, "rb") as pdf_file:
                    st.download_button(
                        label="📥 Baixar PDF",
                        data=pdf_file,
                        file_name="documento_preenchido.pdf",
                        mime="application/pdf"
                    )
    else:
        st.warning("Nenhum campo {{campo}} encontrado.")

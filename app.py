import streamlit as st
from docx import Document
import tempfile
import os
import re
import subprocess

# Extrai campos {{campo}} do docx (inclusive em tabelas)
def extrair_campos(doc):
    campos = set()

    for p in doc.paragraphs:
        campos.update(re.findall(r"\{\{(.*?)\}\}", p.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    campos.update(re.findall(r"\{\{(.*?)\}\}", p.text))

    return list(campos)

# Substitui os campos pelos valores preenchidos
def preencher_campos(doc, dados):
    for p in doc.paragraphs:
        for k, v in dados.items():
            p.text = p.text.replace(f"{{{{{k}}}}}", v)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in dados.items():
                        p.text = p.text.replace(f"{{{{{k}}}}}", v)

# Converte para PDF usando o LibreOffice headless
def converter_para_pdf(caminho_docx):
    saida_dir = tempfile.mkdtemp()
    comando = [
        "soffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", saida_dir,
        caminho_docx
    ]
    subprocess.run(comando, check=True)
    nome_pdf = os.path.splitext(os.path.basename(caminho_docx))[0] + ".pdf"
    return os.path.join(saida_dir, nome_pdf)

# Interface Streamlit
st.title("📄 Gerar PDF fiel ao Word com Docker + LibreOffice")

arquivo = st.file_uploader("Envie o .docx com campos {{campo}}", type="docx")

if arquivo:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(arquivo.read())
        caminho_docx = tmp.name

    doc = Document(caminho_docx)
    campos = extrair_campos(doc)

    if campos:
        st.success("Campos encontrados:")
        dados = {}
        for campo in campos:
            dados[campo] = st.text_input(campo)

        if st.button("Gerar PDF"):
            preencher_campos(doc, dados)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as docx_preenchido:
                doc.save(docx_preenchido.name)
                pdf_path = converter_para_pdf(docx_preenchido.name)

                with open(pdf_path, "rb") as pdf_file:
                    st.download_button(
                        label="📥 Baixar PDF",
                        data=pdf_file,
                        file_name="documento_preenchido.pdf",
                        mime="application/pdf"
                    )
    else:
        st.warning("Nenhum campo {{campo}} encontrado.")
